import re
import os
import markdown
from docx import Document
from docx.shared import Inches
from htmldocx import HtmlToDocx

def main():
    with open("manuscript_draft.md", "r", encoding="utf-8") as f:
        md_text = f.read()
    
    # Map figure placeholders to actual paths
    figures_dir = os.path.join(os.getcwd(), "figures")
    
    # Let's replace the placeholder text with actual image HTML tags so HtmlToDocx handles them
    # But wait, HtmlToDocx might not handle local file images perfectly if not configured, or it might.
    # Alternatively, we can just split the markdown by figures, render HTML for text, and add pictures manually to python-docx Document.
    
    doc = Document()
    new_parser = HtmlToDocx()
    
    # Let's do a line-by-line processing for simplicity to catch the Figure tags, or just block by block.
    blocks = re.split(r'(\*\*Figure \d+:\*\*.*?\n)', md_text)
    
    figure_map = {
        "**Figure 1:**": os.path.join(figures_dir, "Figure_1_Study_Area.png"),
        "**Figure 2:**": os.path.join(figures_dir, "descriptive_stats", "Fig2_TimeSeries_Annual.png"),
        "**Figure 3:**": os.path.join(figures_dir, "trend_analysis", "MK_Significance_NO2.png"),
        "**Figure 4:**": os.path.join(figures_dir, "stl_decomposition", "Fig4a_Seasonal_NO2.png"),
        "**Figure 5:**": os.path.join(figures_dir, "breakpoint_analysis", "Breakpoint_NO2_Terai.png"),
        "**Figure 6:**": os.path.join(figures_dir, "covid_analysis", "Fig6_COVID_Period_Bar.png"),
        "**Figure 7:**": os.path.join(figures_dir, "validation", "Validation_NO2.png")
    }

    for block in blocks:
        if block.startswith("**Figure"):
            # Identify which figure it is
            fig_img = None
            for key, path in figure_map.items():
                if block.startswith(key):
                    fig_img = path
                    break
            
            if fig_img and os.path.exists(fig_img):
                try:
                    doc.add_picture(fig_img, width=Inches(6.0))
                except Exception as e:
                    print(f"Failed to add image {fig_img}: {e}")
            
            # Clean up the text for caption
            caption = block.strip()
            # Remove the bracketed placeholder text since we added the image
            caption = re.sub(r'\[.*?see.*?\]', '', caption)
            doc.add_paragraph(caption)
        else:
            if not block.strip():
                continue
            html = markdown.markdown(block, extensions=['tables'])
            new_parser.add_html_to_document(html, doc)
            
    doc.save("Nepal_Air_Quality_Manuscript.docx")
    print("Successfully generated Nepal_Air_Quality_Manuscript.docx")

if __name__ == "__main__":
    main()
